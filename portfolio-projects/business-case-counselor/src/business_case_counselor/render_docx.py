from __future__ import annotations

from pathlib import Path
from typing import Iterable

from .evidence import EvidenceItem, evidence_appendix
from .model import BusinessCase
from .review import ReviewResult
from .render_md import val


def _require_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Inches, Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError("DOCX export requires python-docx. Install with: pip install python-docx") from exc
    return Document, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, Inches, Pt, OxmlElement, qn


def render_business_case_docx(case: BusinessCase, evidence: list[EvidenceItem], review: ReviewResult, out_path: Path) -> None:
    Document, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, Inches, Pt, OxmlElement, qn = _require_docx()

    def set_cell_shading(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def style_table(t):
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        for i, row in enumerate(t.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
                if i == 0:
                    set_cell_shading(cell, "EAF2F8")
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True

    def add_table(headers: list[str], rows: list[list[str]]):
        t = doc.add_table(rows=1, cols=len(headers))
        for i,h in enumerate(headers):
            t.rows[0].cells[i].text = h
        if not rows:
            rows = [["Not yet defined."] + [""]*(len(headers)-1)]
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row[:len(headers)]):
                cells[i].text = str(v or "")
        style_table(t)

    def add_bullets(items: Iterable[str]):
        vals = [x for x in items if x]
        if not vals:
            doc.add_paragraph("Not yet defined.", style="List Bullet")
        for item in vals:
            doc.add_paragraph(str(item), style="List Bullet")

    def add_para(text: str):
        doc.add_paragraph(text or "TBD")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.5)
    for s in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[s].font.name = "Calibri"

    title = doc.add_heading(f"Business Case: {case.title}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Decision document for approval, prioritization, funding, or governance review")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(["Field", "Detail"], [["Prepared for", case.prepared_for], ["Prepared by", case.prepared_by], ["Date", case.date], ["Version", case.version], ["Classification", case.classification], ["Audience", case.audience_profile]])

    doc.add_heading("Executive Summary", level=1)
    add_para(f"The Problem: {case.problem_statement or 'TBD'}")
    add_para(f"The Recommendation: {case.recommended_solution or 'TBD'}")
    add_para(f"The Financial Case: {case.financial_summary.get('year_1_investment', 'TBD')} investment; {case.financial_summary.get('annual_hard_benefit', 'TBD')} annual hard benefit; {case.financial_summary.get('payback_period', 'TBD')} payback.")
    add_para(f"The Primary Risk: {val(case.risks[0], 'risk', default='TBD') if case.risks else 'TBD'}")
    add_para(f"The Ask: {case.requested_action or 'TBD'} Decision requested from {case.decision_owner or 'TBD'} by {case.decision_timing or 'TBD'}. {case.consequence_of_delay or ''}")

    doc.add_heading("Decision Request", level=1)
    add_table(["Field", "Detail"], [["Decision needed", case.decision_needed], ["Decision owner", case.decision_owner], ["Requested action", case.requested_action], ["Decision timing", case.decision_timing], ["Consequence of delay", case.consequence_of_delay], ["Maturity score", f"{review.total_score}/{review.max_score} - {review.readiness}"]])

    doc.add_heading("1. Problem Statement & Business Need", level=1)
    doc.add_heading("1.1 Observable Symptoms", level=2); add_bullets(case.observable_symptoms)
    doc.add_heading("1.2 Root Cause Analysis", level=2); add_para(case.root_cause_hypothesis)
    doc.add_heading("1.3 Root-Cause Challenge", level=2); add_para(case.root_cause_challenge)
    doc.add_heading("1.4 Quantified Impact of Status Quo", level=2); add_para(case.status_quo_impact)
    doc.add_heading("1.5 Scope Definition", level=2); doc.add_paragraph("In scope:"); add_bullets(case.scope_in); doc.add_paragraph("Out of scope:"); add_bullets(case.scope_out)
    doc.add_heading("1.6 Urgency Drivers", level=2); add_para(case.urgency_driver)

    doc.add_heading("2. Strategic Alignment", level=1)
    add_table(["Organizational Priority", "Connection to Initiative"], [[val(x,"priority"), val(x,"connection")] for x in case.strategic_alignment])

    doc.add_heading("3. Current State Assessment", level=1)
    add_para(case.current_state)
    doc.add_heading("Baseline Performance Metrics", level=2)
    add_table(["Metric", "Current", "Target", "Gap"], [[val(x,"metric"), val(x,"current"), val(x,"target"), val(x,"gap")] for x in case.baseline_metrics])
    doc.add_heading("Identified Gaps", level=2); add_para(case.gap_analysis)
    doc.add_heading("Evidence Sources & Limitations", level=2); add_para(case.evidence_limitations)

    doc.add_heading("4. Options Analysis", level=1)
    doc.add_heading("Evaluation Criteria", level=2)
    add_table(["Criterion", "Weight", "Rationale"], [[val(x,"criterion"), val(x,"weight"), val(x,"rationale")] for x in case.evaluation_criteria])
    doc.add_heading("Option Summaries", level=2)
    add_table(["Option", "Description", "Cost", "Timeline", "Benefits", "Tradeoffs", "Score / Rationale"], [[o.name, o.description, o.cost, o.timeline, "; ".join(o.benefits), "; ".join(o.tradeoffs), o.score or o.rationale] for o in case.options])

    doc.add_heading("5. Recommended Solution", level=1)
    doc.add_heading("5.1 Recommendation", level=2); add_para(case.recommended_solution)
    doc.add_heading("5.2 Rationale", level=2); add_para(case.rationale)
    doc.add_heading("5.3 People / Process / Technology", level=2)
    add_para(f"People: {case.solution_people or 'TBD'}")
    add_para(f"Process: {case.solution_process or 'TBD'}")
    add_para(f"Technology: {case.solution_technology or 'TBD'}")
    doc.add_heading("5.4 Dependencies & Prerequisites", level=2); add_bullets(case.dependencies)
    doc.add_heading("5.5 Organizational Impact", level=2); add_para(case.organizational_impact)

    doc.add_heading("6. Financial Analysis", level=1)
    doc.add_heading("6.1 Investment Summary", level=2)
    add_table(["Category", "One-Time Cost", "Annual Recurring Cost", "Notes / Assumptions"], [[val(c,"category","area","name"), val(c,"one_time","one_time_cost","estimate"), val(c,"recurring","annual_recurring_cost","type"), val(c,"assumption","notes")] for c in case.costs])
    doc.add_heading("6.2 Benefit Model", level=2)
    doc.add_paragraph("Hard Benefits (included in ROI calculation)")
    add_table(["Benefit", "Annual Value", "Basis / Assumption"], [[val(b,"name"), val(b,"annual_value"), val(b,"basis","evidence","evidence_or_assumption")] for b in case.benefits if val(b,"type").lower() == "hard"])
    doc.add_paragraph("Soft Benefits (not included in ROI headline)")
    add_table(["Benefit", "Value", "Basis / Assumption"], [[val(b,"name"), val(b,"annual_value", default="Not included in ROI"), val(b,"basis","evidence","evidence_or_assumption")] for b in case.benefits if val(b,"type").lower() != "hard"])
    doc.add_heading("6.3 Return Analysis", level=2)
    add_table(["Metric", "Value"], [[k.replace("_", " ").title(), str(v)] for k,v in case.financial_summary.items()])
    doc.add_heading("6.4 Key Assumptions", level=2)
    add_table(["Assumption", "Validation"], [[val(a,"assumption","item","name"), val(a,"validation","mitigation","response")] for a in case.assumptions])
    doc.add_heading("6.5 Sensitivity Analysis", level=2)
    add_table(["Scenario", "Variable", "Impact"], [[val(s,"scenario"), val(s,"variable"), val(s,"impact")] for s in case.sensitivity_scenarios])

    doc.add_heading("7. Risk Assessment", level=1)
    add_table(["Risk", "Likelihood", "Impact", "Severity", "Mitigation", "Residual Risk", "Owner"], [[val(r,"risk","item","name"), val(r,"likelihood"), val(r,"impact"), val(r,"severity"), val(r,"mitigation"), val(r,"residual_risk"), val(r,"owner")] for r in case.risks])
    doc.add_heading("Constraints", level=2)
    add_table(["Constraint", "Impact"], [[val(c,"constraint","item","name"), val(c,"impact","mitigation","response")] for c in case.constraints])

    doc.add_heading("8. Implementation Roadmap", level=1)
    add_table(["Phase", "Milestone", "Timing", "Owner", "Deliverables"], [[val(m,"phase"), val(m,"milestone","name"), val(m,"timing"), val(m,"owner"), val(m,"deliverables","notes")] for m in case.implementation_milestones])
    doc.add_heading("Resource Requirements", level=2)
    add_table(["Phase", "Internal FTEs", "External / Vendor", "Technology", "Dependencies"], [[val(r,"phase"), val(r,"internal_fte"), val(r,"external"), val(r,"technology"), val(r,"dependencies")] for r in case.resource_requirements])
    doc.add_heading("Critical Path Items", level=2); add_bullets(case.critical_path)
    doc.add_heading("90-Day Quick Wins", level=2); add_bullets(case.quick_wins)

    doc.add_heading("9. Governance & Decision Framework", level=1)
    add_para(case.governance)
    doc.add_heading("Decision Authority", level=2)
    add_table(["Decision Type", "Accountable", "Responsible", "Consulted", "Informed"], [[val(d,"decision_type"), val(d,"accountable"), val(d,"responsible"), val(d,"consulted"), val(d,"informed")] for d in case.decision_authority])
    doc.add_heading("Escalation Path", level=2); add_para(case.escalation_path)
    doc.add_heading("Success Metrics", level=2)
    add_table(["Metric", "Baseline", "Target", "Frequency", "Owner"], [[val(m,"metric","name","measure"), val(m,"baseline"), val(m,"target"), val(m,"frequency","cadence","review_cadence"), val(m,"owner")] for m in case.success_measures])
    doc.add_heading("Review Gates", level=2)
    add_table(["Gate", "Timing", "Decision Criteria", "Authority"], [[val(g,"gate"), val(g,"timing"), val(g,"criteria"), val(g,"authority")] for g in case.review_gates])

    doc.add_heading("10. Recommendation & Call to Action", level=1)
    add_para(f"We request approval from {case.decision_owner or 'TBD'} to {case.requested_action or 'TBD'} by {case.decision_timing or 'TBD'}.")
    doc.add_heading("Consequence of Delay", level=2); add_para(case.consequence_of_delay)
    doc.add_heading("Next Steps", level=2); add_bullets(case.next_steps)
    doc.add_heading("Open Questions", level=2); add_bullets(case.open_questions)

    doc.add_heading("Critical Professional Counsel", level=1)
    add_para(f"Score: {review.total_score}/{review.max_score}. Readiness: {review.readiness}")
    add_table(["Dimension", "Score", "Finding", "Counsel", "Evidence Needed"], [[c.name, f"{c.score}/5", c.finding, c.counsel, c.evidence_needed] for c in review.criteria])

    doc.add_heading("Appendix: Evidence Summary", level=1)
    if evidence:
        add_table(["Artifact", "Type", "Summary"], [[e.path, e.kind, e.summary] for e in evidence])
    else:
        add_para("No evidence artifacts were provided.")

    footer = doc.add_paragraph("Sample data, where present, is fictional and for demonstration only.")
    for r in footer.runs:
        r.italic = True
        r.font.size = Pt(8)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

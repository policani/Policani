from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


@dataclass
class EvidenceItem:
    path: str
    kind: str
    summary: str
    extracted_text: str


def load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def load_csv(path: Path, max_rows: int = 40) -> str:
    rows = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            rows.append(" | ".join(row))
            if i >= max_rows:
                rows.append("[truncated]")
                break
    return "\n".join(rows)


def load_json(path: Path) -> str:
    data = json.loads(path.read_text(encoding="utf-8"))
    return json.dumps(data, indent=2, ensure_ascii=False)[:12000]


def load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        return f"[python-docx not installed; could not read {path.name}: {exc}]"
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)[:12000]


def load_xlsx(path: Path, max_rows: int = 60, max_cols: int = 12) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        return f"[openpyxl not installed; could not read {path.name}: {exc}]"
    wb = load_workbook(path, data_only=False, read_only=True)
    parts = []
    for ws in wb.worksheets[:8]:
        parts.append(f"# Sheet: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            values = ["" if v is None else str(v) for v in row[:max_cols]]
            if any(values):
                parts.append(" | ".join(values))
            if i + 1 >= max_rows:
                parts.append("[truncated]")
                break
    return "\n".join(parts)[:15000]


def summarize_text(text: str, limit: int = 700) -> str:
    compact = " ".join(text.split())
    if len(compact) <= limit:
        return compact
    return compact[:limit].rstrip() + "..."


def ingest_evidence(folder: Path) -> list[EvidenceItem]:
    if not folder.exists():
        return []
    items: list[EvidenceItem] = []
    loaders = {
        ".md": load_text,
        ".txt": load_text,
        ".csv": load_csv,
        ".json": load_json,
        ".docx": load_docx,
        ".xlsx": load_xlsx,
    }
    for path in sorted(folder.rglob("*")):
        if not path.is_file() or path.name.startswith("~$"):
            continue
        loader = loaders.get(path.suffix.lower())
        if not loader:
            continue
        try:
            text = loader(path)
        except Exception as exc:  # intentional defensive ingestion
            text = f"[could not read {path.name}: {exc}]"
        items.append(EvidenceItem(
            path=str(path.relative_to(folder)),
            kind=path.suffix.lower().lstrip("."),
            summary=summarize_text(text),
            extracted_text=text,
        ))
    return items


def evidence_appendix(items: Iterable[EvidenceItem]) -> str:
    lines = []
    for item in items:
        lines.append(f"### {item.path}")
        lines.append(f"Type: {item.kind}")
        lines.append("")
        lines.append(item.summary or "No extractable text found.")
        lines.append("")
    return "\n".join(lines).strip()

from __future__ import annotations

# Flat renderer for ChatGPT Projects and Codex. Historical source package is archived.

# ===== docx_utils.py =====

import datetime as dt
import os
import tempfile
import zipfile
from pathlib import Path


FIXED_CORE_DATE = dt.datetime(2000, 1, 1, 0, 0, 0)
FIXED_ZIP_DATE = (2000, 1, 1, 0, 0, 0)


def make_docx_zip_deterministic(path: str | Path) -> None:
    """Rewrite a DOCX zip package with stable member order and timestamps."""
    source = Path(path)
    with zipfile.ZipFile(source, "r") as zin:
        entries = [(info.filename, zin.read(info.filename), info.external_attr) for info in zin.infolist()]

    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=str(source.parent))
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for filename, data, external_attr in sorted(entries, key=lambda item: item[0]):
                info = zipfile.ZipInfo(filename, FIXED_ZIP_DATE)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = external_attr
                zout.writestr(info, data)
        temp_path.replace(source)
    finally:
        if temp_path.exists():
            temp_path.unlink()


# ===== models.py =====

import re
from typing import Any


DocumentModel = dict[str, Any]
InlineRun = dict[str, Any]
Block = dict[str, Any]


SECTION_TYPE_ALIASES = {
    "profile": "profile",
    "qualifications": "qualifications",
    "summary": "summary",
    "professional summary": "summary",
    "experience": "experience",
    "professional experience": "experience",
    "work experience": "experience",
    "selected experience": "experience",
    "projects": "projects",
    "selected projects": "projects",
    "education": "education",
    "education and credentials": "education",
    "certifications": "certifications",
    "certification": "certifications",
    "capabilities": "capabilities",
    "skills": "skills",
    "technical skills": "skills",
    "additional": "additional",
    "additional skills": "additional",
    "additional information": "additional",
}


def section_type_for(title: str) -> str:
    normalized = re.sub(r"\s+", " ", title.strip().lower())
    return SECTION_TYPE_ALIASES.get(normalized, slugify(normalized) or "section")


def slugify(value: str) -> str:
    value = re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")
    return value or "section"


def text_run(text: str, *, bold: bool = False, italic: bool = False, color: str | None = None) -> InlineRun:
    run: InlineRun = {"kind": "text", "text": text}
    if bold:
        run["bold"] = True
    if italic:
        run["italic"] = True
    if color:
        run["color"] = color
    return run


def link_run(text: str, url: str, *, bold: bool = False, italic: bool = False, color: str | None = None) -> InlineRun:
    run: InlineRun = {"kind": "link", "text": text, "url": url}
    if bold:
        run["bold"] = True
    if italic:
        run["italic"] = True
    if color:
        run["color"] = color
    return run


def paragraph_block(runs: list[InlineRun], *, style_hint: str | None = None) -> Block:
    block: Block = {"kind": "paragraph", "runs": runs}
    if style_hint:
        block["style_hint"] = style_hint
    return block


def heading_block(runs: list[InlineRun], *, level: int) -> Block:
    return {"kind": "heading", "level": level, "runs": runs}


def bullet_list_block(items: list[list[InlineRun]] | None = None) -> Block:
    return {"kind": "bullet_list", "items": items or []}


def plain_text_from_runs(runs: list[InlineRun]) -> str:
    return "".join(str(run.get("text", "")) for run in runs)


def normalize_model(model: DocumentModel) -> DocumentModel:
    """Return a model with required collection keys present."""
    model.setdefault("type", "resume")
    model.setdefault("title", "")
    model.setdefault("title_runs", [])
    model.setdefault("contact", [])
    model.setdefault("blocks", [])
    model.setdefault("sections", [])
    for section in model["sections"]:
        section.setdefault("type", section_type_for(section.get("heading", "")))
        section.setdefault("blocks", [])
        section.setdefault("items", [])
        for item in section["items"]:
            item.setdefault("blocks", [])
            item.setdefault("components", {})
    return model

# ===== markdown_parser.py =====

import re
from pathlib import Path
from typing import Iterable



HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
BULLET_RE = re.compile(r"^\s{0,3}[-*]\s+(.+?)\s*$")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
EMPHASIS_RE = re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)")
CHANGE_COLOR = "C00000"
KEYWORD_COLOR = "008000"
CHANGE_SPAN_RE = re.compile(
    r"<span\b"
    r"(?=[^>]*\bdata-change\s*=\s*['\"]customized['\"])"
    r"(?=[^>]*\bstyle\s*=\s*['\"][^'\"]*color\s*:\s*#?C00000\b[^'\"]*['\"])"
    r"[^>]*>(.*?)</span>",
    re.IGNORECASE,
)
KEYWORD_SPAN_RE = re.compile(
    r"<span\b"
    r"(?=[^>]*\bdata-origin\s*=\s*['\"]job-description-keyword['\"])"
    r"(?=[^>]*\bstyle\s*=\s*['\"][^'\"]*color\s*:\s*#?008000\b[^'\"]*['\"])"
    r"[^>]*>(.*?)</span>",
    re.IGNORECASE,
)
DATE_RANGE_RE = re.compile(
    r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s+to\s+"
    r"(?:(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}|Present)\b",
    re.IGNORECASE,
)
RESUME_SECTION_TYPES = {
    "profile",
    "qualifications",
    "summary",
    "experience",
    "projects",
    "education",
    "certifications",
    "capabilities",
    "skills",
    "additional",
}
LINE_PRESERVED_SECTION_TYPES = {
    "education",
    "certifications",
    "capabilities",
    "skills",
    "additional",
}


def parse_markdown_file(path: str | Path, doc_type: str | None = None) -> DocumentModel:
    source = Path(path)
    text = source.read_text(encoding="utf-8-sig")
    return parse_markdown(text, doc_type=doc_type or infer_document_type(text, source.name), source=str(source))


def infer_document_type(text: str, source_name: str | None = None) -> str:
    name = (source_name or "").lower()
    lowered = text.lower()
    early_text = lowered[:1000]
    if "cover" in name or "dear " in early_text or re.search(r"(^|\n)\s*(?:\w+\s+)?hiring team,", early_text):
        return "cover_letter"
    return "resume"


def parse_markdown(text: str, doc_type: str | None = None, source: str | None = None) -> DocumentModel:
    kind = doc_type or infer_document_type(text, source)
    if kind not in {"resume", "cover_letter"}:
        raise ValueError("doc_type must be 'resume' or 'cover_letter'")

    model: DocumentModel = {
        "type": kind,
        "title": "",
        "title_runs": [],
        "contact": [],
        "blocks": [],
        "sections": [],
    }
    if source:
        model["source"] = source

    current_section: dict | None = None
    current_item: dict | None = None
    paragraph_lines: list[str] = []
    active_bullet: dict | None = None

    def target_blocks() -> list[dict]:
        if current_item is not None:
            return current_item["blocks"]
        if current_section is not None:
            return current_section["blocks"]
        return model["blocks"]

    def add_section(title: str) -> dict:
        heading = strip_inline_markup(title)
        section = {
            "type": section_type_for(heading),
            "heading": heading,
            "heading_runs": parse_inline_markdown(title),
            "blocks": [],
            "items": [],
        }
        model["sections"].append(section)
        return section

    def add_paragraph(lines: list[str]) -> None:
        nonlocal current_item
        if not lines:
            return
        if kind == "resume" and is_plain_role_header(lines, current_section):
            current_item = make_resume_item_from_lines(lines)
            current_section["items"].append(current_item)
            return

        if should_preserve_resume_lines(lines, current_section, current_item, kind):
            for line in lines:
                add_normalized_paragraph(line)
            return

        add_normalized_paragraph(" ".join(lines))

    def add_normalized_paragraph(markdown_text: str) -> None:
        normalized = normalize_line(markdown_text)
        if not normalized:
            return
        runs = parse_inline_markdown(normalized)
        if kind == "resume" and model["title"] and current_section is None:
            model["contact"].append({"runs": runs})
        else:
            target_blocks().append(paragraph_block(runs))

    def flush_paragraph() -> None:
        nonlocal paragraph_lines, active_bullet
        if paragraph_lines:
            add_paragraph(paragraph_lines)
            paragraph_lines = []
        active_bullet = None

    def add_bullet(markdown_value: str) -> None:
        nonlocal active_bullet
        blocks = target_blocks()
        if not blocks or blocks[-1].get("kind") != "bullet_list":
            blocks.append(bullet_list_block())
        blocks[-1]["items"].append(parse_inline_markdown(normalize_line(markdown_value)))
        active_bullet = blocks[-1]

    def continue_bullet(markdown_value: str) -> None:
        if active_bullet is None or not active_bullet.get("items"):
            paragraph_lines.append(markdown_value)
            return
        active_bullet["items"][-1].append(text_run(" "))
        active_bullet["items"][-1].extend(parse_inline_markdown(normalize_line(markdown_value)))

    for raw_line in normalize_newlines(text).split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            flush_paragraph()
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            heading_text = normalize_line(heading_match.group(2))
            heading_runs = parse_inline_markdown(heading_text)
            heading_plain = plain_text_from_runs(heading_runs)

            if kind == "resume":
                if level == 1:
                    if not model["title"]:
                        model["title"] = heading_plain
                        model["title_runs"] = heading_runs
                        current_section = None
                        current_item = None
                    else:
                        current_section = add_section(heading_text)
                        current_item = None
                elif level == 2:
                    heading_type = section_type_for(heading_plain)
                    if heading_type in RESUME_SECTION_TYPES:
                        current_section = add_section(heading_text)
                        current_item = None
                    else:
                        if current_section is None:
                            current_section = add_section("Profile")
                        current_item = {
                            "heading": heading_plain,
                            "heading_runs": heading_runs,
                            "components": parse_resume_item_components(heading_plain),
                            "blocks": [],
                        }
                        current_section["items"].append(current_item)
                elif level == 3:
                    if current_section is None:
                        current_section = add_section("Profile")
                    current_item = {
                        "heading": heading_plain,
                        "heading_runs": heading_runs,
                        "components": parse_resume_item_components(heading_plain),
                        "blocks": [],
                    }
                    current_section["items"].append(current_item)
                else:
                    target_blocks().append(heading_block(heading_runs, level=level))
            else:
                if level == 1 and not model["title"]:
                    model["title"] = heading_plain
                    model["title_runs"] = heading_runs
                else:
                    model["blocks"].append(heading_block(heading_runs, level=level))
            continue

        bullet_match = BULLET_RE.match(line)
        if bullet_match:
            flush_paragraph()
            add_bullet(bullet_match.group(1))
            continue

        if active_bullet is not None and raw_line[:1].isspace():
            continue_bullet(line)
            continue

        paragraph_lines.append(line.strip())

    flush_paragraph()
    return normalize_model(model)


def is_plain_role_header(lines: list[str], current_section: dict | None) -> bool:
    if not current_section or current_section.get("type") != "experience":
        return False
    cleaned = [normalize_line(line) for line in lines if normalize_line(line)]
    if len(cleaned) < 2 or len(cleaned) > 4:
        return False
    return bool(DATE_RANGE_RE.search(cleaned[-1]))


def should_preserve_resume_lines(
    lines: list[str],
    current_section: dict | None,
    current_item: dict | None,
    doc_type: str,
) -> bool:
    if doc_type != "resume" or current_item is not None or current_section is None:
        return False
    if current_section.get("type") not in LINE_PRESERVED_SECTION_TYPES:
        return False
    cleaned = [normalize_line(line) for line in lines if normalize_line(line)]
    return len(cleaned) > 1


def make_resume_item_from_lines(lines: list[str]) -> dict:
    cleaned = [normalize_line(line) for line in lines if normalize_line(line)]
    heading = " | ".join(cleaned[:3])
    return {
        "heading": heading,
        "heading_runs": parse_inline_markdown(heading),
        "components": {
            "role": cleaned[0],
            "company": cleaned[1] if len(cleaned) > 1 else "",
            "detail": cleaned[2] if len(cleaned) > 2 else "",
        },
        "blocks": [],
    }


def parse_resume_item_components(heading: str) -> dict[str, str]:
    parts = [part.strip() for part in heading.split("|") if part.strip()]
    if len(parts) >= 3:
        return {
            "role": parts[0],
            "company": " | ".join(parts[1:-1]),
            "detail": parts[-1],
        }
    if len(parts) == 2:
        return {"role": parts[0], "company": parts[1], "detail": ""}
    return {"role": heading, "company": "", "detail": ""}


def parse_inline_markdown(text: str, *, color: str | None = None) -> list[InlineRun]:
    runs: list[InlineRun] = []
    position = 0
    while position < len(text):
        match = find_next_inline_token(text, position)
        if match is None:
            append_text_runs(runs, text[position:], color=color)
            break

        append_text_runs(runs, text[position : match.start()], color=color)
        if match.re is LINK_RE:
            label = strip_inline_markup(match.group(1).strip())
            url = match.group(2).strip()
            if label:
                runs.append(link_run(label, url, color=color))
        else:
            runs.extend(parse_inline_markdown(match.group(1), color=span_color(match)))
        position = match.end()
    return merge_adjacent_text_runs(runs)


def find_next_inline_token(text: str, position: int):
    matches = [
        match
        for match in (
            LINK_RE.search(text, position),
            CHANGE_SPAN_RE.search(text, position),
            KEYWORD_SPAN_RE.search(text, position),
        )
        if match
    ]
    if not matches:
        return None
    return min(matches, key=lambda match: match.start())


def span_color(match) -> str:
    return KEYWORD_COLOR if match.re is KEYWORD_SPAN_RE else CHANGE_COLOR


def append_text_runs(runs: list[InlineRun], text: str, *, color: str | None = None) -> None:
    if not text:
        return
    position = 0
    for match in EMPHASIS_RE.finditer(text):
        append_plain_text(runs, text[position : match.start()], color=color)
        token = match.group(0)
        if token.startswith(("**", "__")):
            runs.append(text_run(token[2:-2], bold=True, color=color))
        else:
            runs.append(text_run(token[1:-1], italic=True, color=color))
        position = match.end()
    append_plain_text(runs, text[position:], color=color)


def append_plain_text(runs: list[InlineRun], text: str, *, color: str | None = None) -> None:
    if text:
        runs.append(text_run(text, color=color))


def merge_adjacent_text_runs(runs: Iterable[InlineRun]) -> list[InlineRun]:
    merged: list[InlineRun] = []
    for run in runs:
        if (
            merged
            and run.get("kind") == "text"
            and merged[-1].get("kind") == "text"
            and bool(run.get("bold")) == bool(merged[-1].get("bold"))
            and bool(run.get("italic")) == bool(merged[-1].get("italic"))
            and run.get("color") == merged[-1].get("color")
        ):
            merged[-1]["text"] += run.get("text", "")
        else:
            merged.append(dict(run))
    return merged


def normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def normalize_line(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def strip_inline_markup(text: str) -> str:
    while CHANGE_SPAN_RE.search(text) or KEYWORD_SPAN_RE.search(text):
        text = CHANGE_SPAN_RE.sub(lambda match: match.group(1), text)
        text = KEYWORD_SPAN_RE.sub(lambda match: match.group(1), text)
    text = LINK_RE.sub(lambda match: match.group(1), text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return normalize_line(text)

# ===== renderer.py =====

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor



DEFAULT_CONFIG_FILE = "renderer-config.json"
HORIZONTAL_MARGIN_INCREASE = 3 / 8
VERTICAL_MARGIN_INCREASE = 2 / 8

STYLE_PROFILES: dict[str, dict[str, Any]] = {
    "normal": {
        "resume_font": 9.0,
        "resume_name_font": 18.0,
        "resume_section_font": 12.0,
        "cover_font": 12.0,
        "section_gap_lines": 1.5,
        "section_after": 3,
        "role_before": 4,
        "experience_role_gap_lines": 1.15,
        "body_after": 2,
        "bullet_after": 1,
        "resume_margin": 0.25,
        "cover_margin": 0.85,
    },
    "compact": {
        "resume_font": 9.0,
        "resume_name_font": 17.0,
        "resume_section_font": 12.0,
        "cover_font": 11.5,
        "section_gap_lines": 1.3,
        "section_after": 2,
        "role_before": 3,
        "experience_role_gap_lines": 1.08,
        "body_after": 1,
        "bullet_after": 0,
        "resume_margin": 0.26,
        "cover_margin": 0.78,
    },
    "tight": {
        "resume_font": 8.7,
        "resume_name_font": 16.0,
        "resume_section_font": 10.8,
        "cover_font": 11.0,
        "section_gap_lines": 1.15,
        "section_after": 1,
        "role_before": 2,
        "experience_role_gap_lines": 1.0,
        "body_after": 0,
        "bullet_after": 0,
        "resume_margin": 0.25,
        "cover_margin": 0.75,
    },
}


def load_renderer_config(config_path: str | Path | None = None) -> dict[str, Any]:
    path = Path(config_path) if config_path else Path(__file__).resolve().with_name(DEFAULT_CONFIG_FILE)
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def apply_renderer_config(config: dict[str, Any]) -> None:
    global HORIZONTAL_MARGIN_INCREASE, VERTICAL_MARGIN_INCREASE
    for profile_name, profile_values in config.get("style_profiles", {}).items():
        if profile_name in STYLE_PROFILES and isinstance(profile_values, dict):
            STYLE_PROFILES[profile_name].update(profile_values)
    margins = config.get("margin_adjustments", {})
    if "horizontal_margin_increase" in margins:
        HORIZONTAL_MARGIN_INCREASE = float(margins["horizontal_margin_increase"])
    if "vertical_margin_increase" in margins:
        VERTICAL_MARGIN_INCREASE = float(margins["vertical_margin_increase"])


def config_get(config: dict[str, Any], path: str, default: Any = None) -> Any:
    value: Any = config
    for part in path.split("."):
        if not isinstance(value, dict) or part not in value:
            return default
        value = value[part]
    return value


def render_document(
    model: DocumentModel,
    output_path: str | Path,
    *,
    template_path: str | Path | None = None,
    density: str = "normal",
) -> Path:
    model = normalize_model(model)
    profile = STYLE_PROFILES[density]
    document = Document(str(template_path)) if template_path else Document()
    clear_document_body(document)
    configure_document(document, model["type"], profile)
    ensure_styles(document, model["type"], profile)
    configure_bullet_numbering(document)

    if model["type"] == "resume":
        render_resume(document, model, profile)
    elif model["type"] == "cover_letter":
        render_cover_letter(document, model)
    else:
        raise ValueError("model type must be 'resume' or 'cover_letter'")

    set_core_properties(document)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    make_docx_zip_deterministic(output)
    return output


def clear_document_body(document: DocxDocument) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def configure_document(document: DocxDocument, doc_type: str, profile: dict[str, Any]) -> None:
    margin = profile["resume_margin"] if doc_type == "resume" else profile["cover_margin"]
    for section in document.sections:
        section.top_margin = Inches(margin + VERTICAL_MARGIN_INCREASE)
        section.bottom_margin = Inches(margin + VERTICAL_MARGIN_INCREASE)
        section.left_margin = Inches(margin + HORIZONTAL_MARGIN_INCREASE)
        section.right_margin = Inches(margin + HORIZONTAL_MARGIN_INCREASE)


def configure_bullet_numbering(document: DocxDocument) -> None:
    """Use a renderer-safe Unicode bullet while retaining Word-native numbering."""
    numbering = document.part.numbering_part.element
    abstract_id = None
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == "1":
            abstract_node = num.find(qn("w:abstractNumId"))
            if abstract_node is not None:
                abstract_id = abstract_node.get(qn("w:val"))
            break
    if abstract_id is None:
        return

    for abstract_num in numbering.findall(qn("w:abstractNum")):
        if abstract_num.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for level in abstract_num.findall(qn("w:lvl")):
            num_format = level.find(qn("w:numFmt"))
            if num_format is not None:
                num_format.set(qn("w:val"), "bullet")
            level_text = level.find(qn("w:lvlText"))
            if level_text is not None:
                level_text.set(qn("w:val"), "\u2022")
            run_properties = level.find(qn("w:rPr"))
            if run_properties is None:
                run_properties = OxmlElement("w:rPr")
                level.append(run_properties)
            run_fonts = run_properties.find(qn("w:rFonts"))
            if run_fonts is None:
                run_fonts = OxmlElement("w:rFonts")
                run_properties.append(run_fonts)
            run_fonts.set(qn("w:ascii"), "Calibri")
            run_fonts.set(qn("w:hAnsi"), "Calibri")
        return


def set_core_properties(document: DocxDocument) -> None:
    props = document.core_properties
    props.author = ""
    props.category = "Deterministic Markdown-to-DOCX"
    props.comments = ""
    props.created = FIXED_CORE_DATE
    props.identifier = ""
    props.keywords = ""
    props.language = "en-US"
    props.last_modified_by = ""
    props.modified = FIXED_CORE_DATE
    props.revision = 1
    props.subject = ""
    props.title = ""
    props.version = "1.0"


def ensure_styles(document: DocxDocument, doc_type: str, profile: dict[str, Any]) -> None:
    font = "Segoe UI"
    base_size = profile["resume_font"] if doc_type == "resume" else profile["cover_font"]
    section_before = line_gap_points(base_size, profile["section_gap_lines"]) if doc_type == "resume" else 0

    normal = document.styles["Normal"]
    normal.font.name = font
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    normal.font.size = Pt(base_size)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(profile["body_after"])

    add_or_update_style(
        document,
        "Resume Name",
        font=font,
        size=profile.get("resume_name_font", 18.0),
        bold=True,
        space_after=0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_or_update_style(
        document,
        "Resume Contact",
        font=font,
        size=9,
        space_after=4,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_or_update_style(
        document,
        "Resume Section",
        font=font,
        size=profile.get("resume_section_font", 12.0),
        bold=True,
        all_caps=True,
        space_before=section_before,
        space_after=profile["section_after"],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_or_update_style(
        document,
        "Resume Role",
        font=font,
        size=base_size,
        bold=True,
        space_before=profile["role_before"],
        space_after=0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_or_update_style(
        document,
        "Resume Company",
        font=font,
        size=base_size,
        italic=True,
        space_after=1,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_or_update_style(
        document,
        "Resume Detail",
        font=font,
        size=max(base_size - 0.2, 8.5),
        space_after=1,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_or_update_style(
        document,
        "Resume Body",
        font=font,
        size=base_size,
        space_after=profile["body_after"],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_or_update_style(
        document,
        "Resume Bullet",
        font=font,
        size=base_size,
        space_after=profile["bullet_after"],
        base_style="List Bullet",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        left_indent=0.25,
        first_line_indent=-0.13,
    )

    add_or_update_style(
        document,
        "Letter Name",
        font=font,
        size=12,
        bold=True,
        space_after=0,
    )
    add_or_update_style(
        document,
        "Letter Section",
        font=font,
        size=base_size,
        bold=True,
        space_before=6,
        space_after=3,
    )
    add_or_update_style(
        document,
        "Letter Body",
        font=font,
        size=base_size,
        space_after=0,
        line_spacing=1.08,
    )
    add_or_update_style(
        document,
        "Letter Spacer",
        font=font,
        size=1,
        space_after=7,
        line_spacing=1.0,
    )
    add_or_update_style(
        document,
        "Letter Bullet",
        font=font,
        size=base_size,
        space_after=2,
        base_style="List Bullet",
        left_indent=0.25,
        first_line_indent=-0.13,
    )


def add_or_update_style(
    document: DocxDocument,
    name: str,
    *,
    font: str,
    size: float,
    bold: bool = False,
    italic: bool = False,
    all_caps: bool = False,
    space_before: float = 0,
    space_after: float = 0,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    line_spacing: float = 1.0,
    base_style: str | None = None,
    left_indent: float | None = None,
    first_line_indent: float | None = None,
) -> None:
    styles = document.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    if base_style:
        try:
            style.base_style = styles[base_style]
        except KeyError:
            pass

    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.all_caps = all_caps

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = line_spacing
    paragraph_format.line_spacing_rule = line_spacing_rule_for(line_spacing)
    if alignment is not None:
        paragraph_format.alignment = alignment
    if left_indent is not None:
        paragraph_format.left_indent = Inches(left_indent)
    if first_line_indent is not None:
        paragraph_format.first_line_indent = Inches(first_line_indent)


def line_spacing_rule_for(line_spacing: float):
    return WD_LINE_SPACING.SINGLE if line_spacing == 1.0 else WD_LINE_SPACING.MULTIPLE


def line_gap_points(base_size: float, line_multiple: float) -> float:
    return round(base_size * line_multiple, 2)


def render_resume(document: DocxDocument, model: DocumentModel, profile: dict[str, Any]) -> None:
    if model.get("title_runs"):
        paragraph = document.add_paragraph(style="Resume Name")
        add_runs(paragraph, model["title_runs"])
    elif model.get("title"):
        document.add_paragraph(str(model["title"]), style="Resume Name")

    for contact_line in model.get("contact", []):
        paragraph = document.add_paragraph(style="Resume Contact")
        add_runs(paragraph, contact_line.get("runs", []))

    for block in model.get("blocks", []):
        render_block(document, block, paragraph_style="Resume Body", bullet_style="Resume Bullet")

    for section in model.get("sections", []):
        paragraph = document.add_paragraph(style="Resume Section")
        paragraph.paragraph_format.keep_with_next = True
        add_runs(paragraph, section.get("heading_runs") or [{"kind": "text", "text": section["heading"]}])
        add_bottom_border(paragraph)

        for block in section.get("blocks", []):
            if is_qualification_capabilities_header(section, block):
                render_resume_subsection_header(document, block, profile)
            else:
                render_block(document, block, paragraph_style="Resume Body", bullet_style="Resume Bullet")

        for item_index, item in enumerate(section.get("items", [])):
            render_resume_item(document, item, profile, section_type=section.get("type", ""), item_index=item_index)


def render_resume_item(
    document: DocxDocument,
    item: dict[str, Any],
    profile: dict[str, Any],
    *,
    section_type: str,
    item_index: int,
) -> None:
    components = item.get("components", {})
    role = components.get("role", "")
    company = components.get("company", "")
    detail = components.get("detail", "")

    if company:
        role_paragraph = document.add_paragraph(style="Resume Role")
        add_runs(role_paragraph, parse_inline_markdown(role))
        apply_resume_item_spacing(role_paragraph, profile, section_type=section_type, item_index=item_index)
        role_paragraph.paragraph_format.keep_with_next = True
        company_paragraph = document.add_paragraph(style="Resume Company")
        add_runs(company_paragraph, parse_inline_markdown(company))
        company_paragraph.paragraph_format.keep_with_next = True
        if detail:
            detail_paragraph = document.add_paragraph(style="Resume Detail")
            add_runs(detail_paragraph, parse_inline_markdown(detail))
            detail_paragraph.paragraph_format.keep_with_next = True
    else:
        paragraph = document.add_paragraph(style="Resume Role")
        apply_resume_item_spacing(paragraph, profile, section_type=section_type, item_index=item_index)
        paragraph.paragraph_format.keep_with_next = True
        add_runs(paragraph, item.get("heading_runs") or [{"kind": "text", "text": item.get("heading", "")}])

    for block in item.get("blocks", []):
        render_block(document, block, paragraph_style="Resume Body", bullet_style="Resume Bullet")


def apply_resume_item_spacing(
    paragraph,
    profile: dict[str, Any],
    *,
    section_type: str,
    item_index: int,
) -> None:
    if section_type == "experience":
        line_multiple = profile["experience_role_gap_lines"] if item_index > 0 else 0
        paragraph.paragraph_format.space_before = Pt(line_gap_points(profile["resume_font"], line_multiple))


def is_qualification_capabilities_header(section: dict[str, Any], block: dict[str, Any]) -> bool:
    if section.get("type") != "qualifications" or block.get("kind") != "paragraph":
        return False
    return plain_text_from_runs(block.get("runs", [])).strip().lower() == "capabilities"


def render_resume_subsection_header(
    document: DocxDocument,
    block: dict[str, Any],
    profile: dict[str, Any],
) -> None:
    paragraph = document.add_paragraph(style="Resume Role")
    paragraph.paragraph_format.space_before = Pt(line_gap_points(profile["resume_font"], 1.0))
    paragraph.paragraph_format.keep_with_next = True
    add_runs(paragraph, block.get("runs", []))


def render_cover_letter(document: DocxDocument, model: DocumentModel) -> None:
    blocks = model.get("blocks", [])
    if model.get("title_runs"):
        paragraph = document.add_paragraph(style="Letter Name")
        add_runs(paragraph, model["title_runs"])
        if blocks:
            add_letter_copy_spacer(document)

    for index, block in enumerate(blocks):
        render_block(document, block, paragraph_style="Letter Body", bullet_style="Letter Bullet")
        if should_add_letter_copy_spacer(blocks, index):
            add_letter_copy_spacer(document)


def should_add_letter_copy_spacer(blocks: list[dict[str, Any]], index: int) -> bool:
    if index >= len(blocks) - 1:
        return False
    return blocks[index].get("kind") in {"paragraph", "bullet_list"}


def add_letter_copy_spacer(document: DocxDocument) -> None:
    document.add_paragraph(style="Letter Spacer")


def render_block(
    document: DocxDocument,
    block: dict[str, Any],
    *,
    paragraph_style: str,
    bullet_style: str,
) -> None:
    kind = block.get("kind")
    if kind == "paragraph":
        paragraph = document.add_paragraph(style=paragraph_style)
        add_runs(paragraph, block.get("runs", []))
    elif kind == "heading":
        paragraph = document.add_paragraph(style="Letter Section" if paragraph_style == "Letter Body" else "Resume Role")
        add_runs(paragraph, block.get("runs", []))
    elif kind == "bullet_list":
        for item_runs in block.get("items", []):
            paragraph = document.add_paragraph(style=bullet_style)
            apply_bullet_numbering(paragraph)
            add_runs(paragraph, item_runs)


def add_runs(paragraph, runs: list[InlineRun]) -> None:
    for run in runs:
        text = str(run.get("text", ""))
        if not text:
            continue
        if run.get("kind") == "link":
            add_hyperlink(
                paragraph,
                text,
                str(run.get("url", "")),
                bold=bool(run.get("bold")),
                italic=bool(run.get("italic")),
                color=run.get("color"),
            )
        else:
            word_run = paragraph.add_run(text)
            if run.get("bold"):
                word_run.bold = True
            if run.get("italic"):
                word_run.italic = True
            if run.get("color"):
                word_run.font.color.rgb = RGBColor.from_string(str(run["color"]))


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color or "0563C1")
    run_properties.append(color_node)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    if bold:
        run_properties.append(OxmlElement("w:b"))
    if italic:
        run_properties.append(OxmlElement("w:i"))

    text_element = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text

    new_run.append(run_properties)
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def apply_bullet_numbering(paragraph, *, num_id: int = 1, level: int = 0) -> None:
    """Apply explicit Word numbering so custom bullet styles render consistently."""
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_pr = paragraph_properties.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        paragraph_properties.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))

    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_bottom_border(paragraph) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    border = paragraph_properties.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        paragraph_properties.append(border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A6A6A6")
    border.append(bottom)


def estimate_word_count(model: DocumentModel) -> int:
    words: list[str] = []

    def add_runs_to_words(runs: list[InlineRun]) -> None:
        words.extend(plain_text_from_runs(runs).split())

    add_runs_to_words(model.get("title_runs", []))
    for contact in model.get("contact", []):
        add_runs_to_words(contact.get("runs", []))
    for block in model.get("blocks", []):
        collect_block_words(block, add_runs_to_words)
    for section in model.get("sections", []):
        add_runs_to_words(section.get("heading_runs", []))
        for block in section.get("blocks", []):
            collect_block_words(block, add_runs_to_words)
        for item in section.get("items", []):
            add_runs_to_words(item.get("heading_runs", []))
            for block in item.get("blocks", []):
                collect_block_words(block, add_runs_to_words)
    return len(words)


def collect_block_words(block: dict[str, Any], collector) -> None:
    if block.get("kind") in {"paragraph", "heading"}:
        collector(block.get("runs", []))
    elif block.get("kind") == "bullet_list":
        for item in block.get("items", []):
            collector(item)

# ===== template.py =====

from pathlib import Path

from docx import Document



def create_reference_docx(output_path: str | Path) -> Path:
    """Create a base reference.docx containing all renderer styles."""
    document = Document()
    clear_document_body(document)
    ensure_styles(document, "resume", STYLE_PROFILES["normal"])
    ensure_styles(document, "cover_letter", STYLE_PROFILES["normal"])
    configure_bullet_numbering(document)
    set_core_properties(document)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    make_docx_zip_deterministic(output)
    return output

# ===== validation.py =====

import re
import shutil
import subprocess
import sys
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from lxml import etree
from docx.oxml.ns import qn



NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}

RAW_URL_RE = re.compile(r"\b(?:https?://|www\.)\S+", re.IGNORECASE)
MARKDOWN_LINK_RE = re.compile(r"\[[^\]]+\]\([^)]+\)")
MARKDOWN_BULLET_RE = re.compile(r"^\s*[-*]\s+")
REVIEW_SPAN_TAG_RE = re.compile(r"</?span\b", re.IGNORECASE)


@dataclass
class ValidationIssue:
    severity: str
    code: str
    message: str


@dataclass
class ValidationResult:
    issues: list[ValidationIssue] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def ok(self) -> bool:
        return not any(issue.severity == "error" for issue in self.issues)

    def add(self, severity: str, code: str, message: str) -> None:
        self.issues.append(ValidationIssue(severity, code, message))


@dataclass
class DocxInspection:
    paragraphs: list[dict[str, str]]
    text: str
    table_count: int
    drawing_count: int
    text_box_count: int
    hyperlink_count: int
    manual_page_break_count: int
    relationship_hyperlinks: list[dict[str, str]]
    styles: set[str]
    column_counts: list[int]


def validate_model(model: DocumentModel) -> ValidationResult:
    result = ValidationResult()
    doc_type = model.get("type")
    if doc_type not in {"resume", "cover_letter"}:
        result.add("error", "MODEL_TYPE", "Model type must be 'resume' or 'cover_letter'.")

    if doc_type == "resume":
        if not model.get("title"):
            result.add("error", "RESUME_TITLE", "Resume model must include a candidate name as title.")
        if not model.get("sections"):
            result.add("error", "RESUME_SECTIONS", "Resume model must include at least one section.")
    if doc_type == "cover_letter":
        words = estimate_word_count(model)
        result.metadata["word_count"] = words
        if words and not 250 <= words <= 350:
            result.add("warning", "COVER_WORD_COUNT", f"Cover letter word count is {words}; preferred range is 250-350.")

    for index, section in enumerate(model.get("sections", []), start=1):
        if not section.get("heading"):
            result.add("error", "SECTION_HEADING", f"Section {index} is missing a heading.")

    return result


def validate_docx(
    path: str | Path,
    *,
    doc_type: str = "resume",
    max_pages: int | None = None,
    rendered_pages_dir: str | Path | None = None,
    expected_hyperlinks: int | None = None,
) -> ValidationResult:
    result = ValidationResult()
    inspection = inspect_docx(path)
    result.metadata.update(
        {
            "paragraph_count": len(inspection.paragraphs),
            "table_count": inspection.table_count,
            "hyperlink_count": inspection.hyperlink_count,
        }
    )

    for paragraph in inspection.paragraphs:
        text = paragraph["text"]
        if paragraph.get("style") in {"ResumeBullet", "LetterBullet"} and paragraph.get("has_num_pr") != "1":
            result.add("error", "BULLET_NUMPR", f"Bullet paragraph is missing Word numbering properties: {text[:80]}")
        if MARKDOWN_BULLET_RE.match(text):
            result.add("error", "LITERAL_BULLET", f"Literal Markdown bullet remains visible: {text[:80]}")
        if MARKDOWN_LINK_RE.search(text):
            result.add("error", "MARKDOWN_LINK", f"Markdown link syntax remains visible: {text[:80]}")
        if REVIEW_SPAN_TAG_RE.search(text):
            result.add("error", "REVIEW_SPAN_TAG", f"Review span markup remains visible: {text[:80]}")
        hyperlink_texts = set(paragraph.get("hyperlink_texts", []))
        for raw_url in RAW_URL_RE.finditer(text):
            visible_url = raw_url.group(0)
            if visible_url not in hyperlink_texts:
                result.add("error", "RAW_URL", f"Raw URL is visible but not clickable: {visible_url}")

    if doc_type == "resume":
        if inspection.table_count:
            result.add("error", "RESUME_TABLES", "Resume output contains tables, which breaks ATS single-column rules.")
        if inspection.text_box_count:
            result.add("error", "RESUME_TEXT_BOXES", "Resume output contains text boxes.")
        if inspection.drawing_count:
            result.add("error", "RESUME_DRAWINGS", "Resume output contains drawings or images.")
        if any(count > 1 for count in inspection.column_counts):
            result.add("error", "RESUME_COLUMNS", "Resume output contains multi-column section settings.")
        if inspection.manual_page_break_count:
            result.add("error", "RESUME_MANUAL_PAGE_BREAK", "Resume output contains hidden manual page-break markup.")
        required_styles = {"Resume Name", "Resume Section", "Resume Role", "Resume Company", "Resume Body", "Resume Bullet"}
        missing_styles = sorted(required_styles - inspection.styles)
        if missing_styles:
            result.add("error", "MISSING_STYLES", f"Missing required resume styles: {', '.join(missing_styles)}")

    if expected_hyperlinks is not None and inspection.hyperlink_count != expected_hyperlinks:
        result.add(
            "error",
            "HYPERLINK_COUNT",
            f"Expected {expected_hyperlinks} hyperlinks, found {inspection.hyperlink_count}.",
        )

    for relationship in inspection.relationship_hyperlinks:
        if not relationship.get("target"):
            result.add("error", "HYPERLINK_TARGET", "Hyperlink relationship is missing a target URL.")

    if max_pages is not None:
        page_count = count_rendered_pages(rendered_pages_dir) if rendered_pages_dir else None
        result.metadata["page_count"] = page_count
        if page_count is None:
            result.add(
                "warning",
                "PAGE_COUNT_NOT_RENDERED",
                "Page count was requested but no rendered page image directory was supplied.",
            )
        elif page_count > max_pages:
            result.add("error", "PAGE_LIMIT", f"Rendered page count is {page_count}; max allowed is {max_pages}.")

    return result


def inspect_docx(path: str | Path) -> DocxInspection:
    docx_path = Path(path)
    with zipfile.ZipFile(docx_path, "r") as package:
        document_xml = etree.fromstring(package.read("word/document.xml"))
        styles_xml = etree.fromstring(package.read("word/styles.xml"))
        rels_xml = None
        rels_name = "word/_rels/document.xml.rels"
        if rels_name in package.namelist():
            rels_xml = etree.fromstring(package.read(rels_name))

    paragraphs: list[dict[str, str]] = []
    for paragraph in document_xml.xpath(".//w:p", namespaces=NS):
        text = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
        hyperlink_texts = [
            "".join(hyperlink.xpath(".//w:t/text()", namespaces=NS))
            for hyperlink in paragraph.xpath(".//w:hyperlink", namespaces=NS)
        ]
        style_node = paragraph.find("w:pPr/w:pStyle", namespaces=NS)
        num_pr = paragraph.find("w:pPr/w:numPr", namespaces=NS)
        page_break_before = paragraph.find("w:pPr/w:pageBreakBefore", namespaces=NS)
        page_breaks = paragraph.xpath(".//w:br[@w:type='page']", namespaces=NS)
        style = style_node.get(qn("w:val")) if style_node is not None else ""
        paragraphs.append(
            {
                "text": text,
                "style": style,
                "has_num_pr": "1" if num_pr is not None else "0",
                "page_break_before": "1" if page_break_before is not None else "0",
                "manual_page_breaks": str(len(page_breaks)),
                "hyperlink_texts": hyperlink_texts,
            }
        )

    style_names = {
        node.get(qn("w:val"), "")
        for node in styles_xml.xpath(".//w:style/w:name", namespaces=NS)
        if node.get(qn("w:val"))
    }
    column_counts = [
        int(node.get(qn("w:num"), "1"))
        for node in document_xml.xpath(".//w:sectPr/w:cols", namespaces=NS)
        if node.get(qn("w:num"), "1").isdigit()
    ]
    relationship_hyperlinks: list[dict[str, str]] = []
    if rels_xml is not None:
        for rel in rels_xml.xpath(".//rel:Relationship", namespaces=NS):
            rel_type = rel.get("Type", "")
            if rel_type.endswith("/hyperlink"):
                relationship_hyperlinks.append(
                    {
                        "id": rel.get("Id", ""),
                        "target": rel.get("Target", ""),
                        "target_mode": rel.get("TargetMode", ""),
                    }
                )

    return DocxInspection(
        paragraphs=paragraphs,
        text="\n".join(paragraph["text"] for paragraph in paragraphs),
        table_count=len(document_xml.xpath(".//w:tbl", namespaces=NS)),
        drawing_count=len(document_xml.xpath(".//w:drawing", namespaces=NS)),
        text_box_count=len(document_xml.xpath(".//w:txbxContent", namespaces=NS)),
        hyperlink_count=len(document_xml.xpath(".//w:hyperlink", namespaces=NS)),
        manual_page_break_count=len(document_xml.xpath(".//w:pPr/w:pageBreakBefore | .//w:br[@w:type='page']", namespaces=NS)),
        relationship_hyperlinks=relationship_hyperlinks,
        styles=style_names,
        column_counts=column_counts,
    )


def extract_plain_text(path: str | Path) -> str:
    return inspect_docx(path).text


def count_model_hyperlinks(model: DocumentModel) -> int:
    count = 0

    def count_runs(runs: list[dict[str, Any]]) -> None:
        nonlocal count
        count += sum(1 for run in runs if run.get("kind") == "link")

    count_runs(model.get("title_runs", []))
    for contact in model.get("contact", []):
        count_runs(contact.get("runs", []))
    for block in model.get("blocks", []):
        count_block_hyperlinks(block, count_runs)
    for section in model.get("sections", []):
        count_runs(section.get("heading_runs", []))
        for block in section.get("blocks", []):
            count_block_hyperlinks(block, count_runs)
        for item in section.get("items", []):
            count_runs(item.get("heading_runs", []))
            for block in item.get("blocks", []):
                count_block_hyperlinks(block, count_runs)
    return count


def count_block_hyperlinks(block: dict[str, Any], collector) -> None:
    if block.get("kind") in {"paragraph", "heading"}:
        collector(block.get("runs", []))
    elif block.get("kind") == "bullet_list":
        for item in block.get("items", []):
            collector(item)


def count_rendered_pages(rendered_pages_dir: str | Path | None) -> int | None:
    if not rendered_pages_dir:
        return None
    directory = Path(rendered_pages_dir)
    if not directory.exists():
        return None
    pages = sorted(directory.glob("page-*.png"))
    return len(pages) if pages else None


def render_docx_pages(
    docx_path: str | Path,
    output_dir: str | Path,
    *,
    render_script: str | Path | None = None,
    python_executable: str | Path | None = None,
) -> int:
    script = resolve_render_script(render_script)
    out = Path(output_dir)
    if out.exists():
        shutil.rmtree(out)
    out.mkdir(parents=True, exist_ok=True)

    if sys.platform == "win32":
        return render_docx_pages_with_word(docx_path, out)

    if script is None:
        raise FileNotFoundError("Could not find render_docx.py. Pass --render-script or set DOCX_RENDER_SCRIPT.")
    python = str(python_executable or sys.executable)
    subprocess.run(
        [
            python,
            str(script),
            str(docx_path),
            "--output_dir",
            str(out),
        ],
        check=True,
    )
    page_count = count_rendered_pages(out)
    if page_count is None:
        raise RuntimeError(f"Renderer produced no page images in {out}")
    return page_count


def render_docx_pages_with_word(docx_path: str | Path, output_dir: str | Path) -> int:
    try:
        import pypdfium2 as pdfium
    except Exception as exc:
        raise RuntimeError("pypdfium2 is required for Windows page rendering.") from exc

    docx = Path(docx_path).resolve()
    out = Path(output_dir).resolve()
    out.mkdir(parents=True, exist_ok=True)
    pdf_path = out / "word-export.pdf"

    with tempfile.NamedTemporaryFile(prefix=f"md2docx_word_export_{os.getpid()}_", suffix=".ps1", delete=False) as script_file:
        script_path = Path(script_file.name)
    escaped_docx = str(docx).replace("'", "''")
    escaped_pdf = str(pdf_path).replace("'", "''")
    script_path.write_text(
        "\n".join(
            [
                f"$docx = '{escaped_docx}'",
                f"$pdf = '{escaped_pdf}'",
                "$word = $null",
                "$doc = $null",
                "try {",
                "    $word = New-Object -ComObject Word.Application",
                "    $word.Visible = $false",
                "    $word.DisplayAlerts = 0",
                "    $doc = $word.Documents.Open($docx)",
                "    $doc.ExportAsFixedFormat($pdf, 17)",
                "    $doc.Close($false)",
                "    $word.Quit()",
                "}",
                "catch {",
                "    if ($null -ne $doc) { $doc.Close($false) }",
                "    if ($null -ne $word) { $word.Quit() }",
                "    Write-Error $_.Exception.Message",
                "    exit 1",
                "}",
            ]
        ),
        encoding="utf-8",
    )
    try:
        subprocess.run(
            [
                "powershell.exe",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=90,
        )
    finally:
        script_path.unlink(missing_ok=True)

    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise RuntimeError(f"Microsoft Word did not create a nonempty PDF: {pdf_path}")

    pdf = pdfium.PdfDocument(str(pdf_path))
    for index, page in enumerate(pdf, start=1):
        image = page.render(scale=2).to_pil()
        image.save(out / f"page-{index}.png")

    page_count = count_rendered_pages(out)
    if page_count is None:
        raise RuntimeError(f"Microsoft Word renderer produced no page images in {out}")
    return page_count


def resolve_render_script(render_script: str | Path | None = None) -> Path | None:
    if render_script:
        path = Path(render_script)
        return path if path.exists() else None
    import os

    env_value = os.environ.get("DOCX_RENDER_SCRIPT")
    if env_value:
        env_path = Path(env_value)
        if env_path.exists():
            return env_path
    for parent in Path.home().glob(".codex/plugins/cache/openai-primary-runtime/documents/*/skills/documents/render_docx.py"):
        if parent.exists():
            return parent
    return None


def print_validation_result(result: ValidationResult) -> None:
    if result.metadata:
        for key, value in result.metadata.items():
            print(f"{key}: {value}")
    if not result.issues:
        print("Validation passed.")
        return
    for issue in result.issues:
        print(f"{issue.severity.upper()} {issue.code}: {issue.message}")

# ===== cli.py =====

import argparse
import hashlib
import json
from pathlib import Path
from typing import Any


VALIDATION_RULES_VERSION = "2026-04-29.1"
RESULT_PASS = "pass"
RESULT_PASS_WITH_VISUAL_WARNING = "pass_with_visual_qa_warning"
RESULT_FAIL_PAGE_COUNT = "fail_page_count"
RESULT_FAIL_STRUCTURE = "fail_structure"
RESULT_FAIL_RENDERER = "fail_renderer"
RESULT_FAIL_HYPERLINKS = "fail_hyperlinks"


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        return args.func(args)
    except Exception as exc:
        print(f"ERROR: {exc}")
        return 1


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(prog="md2docx-resume")
    subparsers = parser.add_subparsers(required=True)

    template = subparsers.add_parser("template", help="Create reference.docx with renderer styles.")
    template.add_argument("-o", "--output", default="reference.docx")
    template.add_argument("--config")
    template.set_defaults(func=cmd_template)

    to_json = subparsers.add_parser("to-json", help="Convert Markdown to the intermediate JSON model.")
    to_json.add_argument("input")
    to_json.add_argument("--type", choices=["resume", "cover_letter"], dest="doc_type")
    to_json.add_argument("-o", "--output", required=True)
    to_json.set_defaults(func=cmd_to_json)

    render = subparsers.add_parser("render", help="Render Markdown or JSON to DOCX.")
    render.add_argument("input")
    render.add_argument("--type", choices=["resume", "cover_letter"], dest="doc_type")
    render.add_argument("-o", "--output")
    render.add_argument("--template")
    render.add_argument("--config")
    render.add_argument("--mode", choices=["fast", "qa", "debug"])
    render.add_argument("--density")
    render.add_argument("--max-pages", type=int)
    render.add_argument("--render-script")
    render.add_argument("--render-dir")
    render.add_argument("--qa-report")
    render.add_argument("--download-dir")
    render.add_argument("--cache-file")
    render.add_argument("--cache-validation", choices=["true", "false"])
    render.add_argument("--force-visual-qa", action="store_true")
    render.add_argument("--preserve-intermediates", action="store_true")
    render.add_argument("--no-validate", action="store_true")
    render.set_defaults(func=cmd_render)

    validate = subparsers.add_parser("validate", help="Validate a rendered DOCX.")
    validate.add_argument("input")
    validate.add_argument("--type", choices=["resume", "cover_letter"], default="resume", dest="doc_type")
    validate.add_argument("--max-pages", type=int)
    validate.add_argument("--rendered-pages")
    validate.set_defaults(func=cmd_validate)

    extract = subparsers.add_parser("extract-text", help="Extract plain text from DOCX.")
    extract.add_argument("input")
    extract.set_defaults(func=cmd_extract_text)

    package = subparsers.add_parser("package-downloads", help="Bundle multiple rendered DOCX files into one ZIP download artifact.")
    package.add_argument("inputs", nargs="+")
    package.add_argument("-o", "--output")
    package.add_argument("--download-dir")
    package.set_defaults(func=cmd_package_downloads)

    return parser


def cmd_template(args: argparse.Namespace) -> int:
    config = load_renderer_config(args.config)
    apply_renderer_config(config)
    output = create_reference_docx(args.output)
    print(f"Created {output}")
    return 0


def cmd_to_json(args: argparse.Namespace) -> int:
    model = parse_markdown_file(args.input, doc_type=args.doc_type)
    result = validate_model(model)
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(model, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(f"Wrote {output}")
    print_validation_result(result)
    return 0 if result.ok else 1


def cmd_render(args: argparse.Namespace) -> int:
    config = load_renderer_config(args.config)
    apply_renderer_config(config)
    apply_render_defaults(args, config)
    model = load_input_model(args.input, args.doc_type)
    model_result = validate_model(model)
    if not model_result.ok:
        print_validation_result(model_result)
        return 1

    output = resolve_output_path(args.output, args.input, config)
    cache_enabled = args.cache_validation == "true"
    cache_file = Path(args.cache_file)
    fingerprint = build_fingerprint(args.template, args.config)
    effective_mode = resolve_effective_mode(args.mode, cache_enabled, cache_file, model["type"], fingerprint)
    if args.force_visual_qa and effective_mode == "fast":
        effective_mode = "qa"

    visual_qa_requested = effective_mode in {"qa", "debug"}
    rendered_dir = resolve_render_dir(args.render_dir, output) if visual_qa_requested else None
    densities = density_attempts(args.density, model["type"], visual_qa_requested)
    final_result = None
    final_density = densities[-1]
    final_page_count = None
    visual_qa_performed = False
    visual_qa_reason = "Skipped in fast mode because structural validation is the default production check."
    renderer_error = ""
    trim_report_path = None
    report_path = Path(args.qa_report) if args.qa_report else None

    for density in densities:
        target_output = output
        if effective_mode == "debug" or args.preserve_intermediates:
            target_output = output.with_name(f"{output.stem}.{density}{output.suffix}")
        render_document(model, target_output, template_path=args.template, density=density)
        if target_output != output:
            render_document(model, output, template_path=args.template, density=density)
        page_count = None
        rendered_for_validation = None
        if rendered_dir:
            page_dir = rendered_dir / density if len(densities) > 1 else rendered_dir
            try:
                page_count = render_docx_pages(output, page_dir, render_script=args.render_script)
                rendered_for_validation = page_dir
                visual_qa_performed = True
                visual_qa_reason = f"Performed because mode={effective_mode}."
            except Exception as exc:
                renderer_error = str(exc)
                visual_qa_reason = f"Renderer failed or was unavailable: {renderer_error}"

        if args.no_validate:
            print(f"Rendered {output} using density={density}")
            return 0

        final_result = validate_docx(
            output,
            doc_type=model["type"],
            max_pages=args.max_pages if rendered_for_validation else None,
            rendered_pages_dir=rendered_for_validation,
            expected_hyperlinks=count_model_hyperlinks(model),
        )
        final_density = density
        final_page_count = page_count
        final_result.metadata["density"] = density
        final_result.metadata["mode"] = effective_mode
        if page_count is not None:
            final_result.metadata["page_count"] = page_count
        if final_result.ok:
            download_artifact = maybe_prepare_download_artifact(output, args.download_dir)
            if visual_qa_performed and cache_enabled:
                update_validation_cache(cache_file, model["type"], fingerprint, density)
            report = build_qa_report(
                args=args,
                model=model,
                output=output,
                result=final_result,
                fingerprint=fingerprint,
                mode=effective_mode,
                density=density,
                page_count=page_count,
                visual_qa_performed=visual_qa_performed,
                visual_qa_reason=visual_qa_reason,
                renderer_error=renderer_error,
                download_artifact=download_artifact,
            )
            write_qa_report(report_path, output, report)
            print(f"Rendered {output}")
            if download_artifact:
                print(f"Download artifact: {download_artifact}")
            print_validation_result(final_result)
            return 0
        if renderer_error:
            break
        if not has_issue(final_result, "PAGE_LIMIT"):
            break

    assert final_result is not None
    if has_issue(final_result, "PAGE_LIMIT") and model["type"] == "resume":
        trim_report_path = write_trim_required_report(output, model, final_result, final_density, args.max_pages)
        final_result.metadata["trim_report"] = str(trim_report_path)
    report = build_qa_report(
        args=args,
        model=model,
        output=output,
        result=final_result,
        fingerprint=fingerprint,
        mode=effective_mode,
        density=final_density,
        page_count=final_page_count,
        visual_qa_performed=visual_qa_performed,
        visual_qa_reason=visual_qa_reason,
        renderer_error=renderer_error,
        trim_report_path=trim_report_path,
    )
    write_qa_report(report_path, output, report)
    if renderer_error and final_result.ok:
        print(f"Rendered {output}, but visual QA failed.")
        print_validation_result(final_result)
        return 0
    print(f"Rendered {output}, but validation failed.")
    print_validation_result(final_result)
    return 1


def cmd_validate(args: argparse.Namespace) -> int:
    result = validate_docx(
        args.input,
        doc_type=args.doc_type,
        max_pages=args.max_pages,
        rendered_pages_dir=args.rendered_pages,
    )
    print_validation_result(result)
    return 0 if result.ok else 1


def cmd_extract_text(args: argparse.Namespace) -> int:
    print(extract_plain_text(args.input))
    return 0


def cmd_package_downloads(args: argparse.Namespace) -> int:
    artifacts = [Path(value) for value in args.inputs]
    if len(artifacts) < 2:
        raise ValueError("package-downloads requires at least two input DOCX files.")
    zip_artifact = prepare_download_bundle(artifacts, download_dir=args.download_dir or "/mnt/data", output_name=args.output)
    print(f"Download artifact: {zip_artifact}")
    return 0


def load_input_model(path: str, doc_type: str | None = None) -> dict[str, Any]:
    source = Path(path)
    if source.suffix.lower() == ".json":
        return json.loads(source.read_text(encoding="utf-8"))
    return parse_markdown_file(source, doc_type=doc_type)


def apply_render_defaults(args: argparse.Namespace, config: dict[str, Any]) -> None:
    args.mode = args.mode or config_get(config, "render_defaults.mode", "fast")
    args.density = args.density or config_get(config, "render_defaults.density", "auto")
    if args.density != "auto" and args.density not in STYLE_PROFILES:
        raise ValueError(f"Unknown density '{args.density}'. Use auto or one of: {', '.join(STYLE_PROFILES)}")
    if args.template is None:
        template = config_get(config, "paths.template", "")
        template_path = Path(template) if template else None
        if template_path and not template_path.exists():
            template_path = Path(__file__).resolve().with_name(template)
        args.template = str(template_path) if template_path and template_path.exists() else None
    if args.cache_file is None:
        args.cache_file = config_get(config, "paths.cache_file", "outputs/render-validation-cache.json")
    if args.cache_validation is None:
        enabled = bool(config_get(config, "render_defaults.cache_validation", True))
        args.cache_validation = "true" if enabled else "false"
    if args.download_dir is None:
        configured = str(config_get(config, "paths.download_dir", "") or "").strip()
        args.download_dir = configured or None


def resolve_output_path(output: str | None, input_path: str, config: dict[str, Any]) -> Path:
    if output:
        return Path(output)
    output_dir = str(config_get(config, "paths.default_output_dir", "outputs") or ".")
    return Path(output_dir) / (Path(input_path).stem + ".docx")


def slugify_artifact_filename(filename: str) -> str:
    path = Path(filename)
    suffix = path.suffix.lower()
    safe_stem = re.sub(r"[^A-Za-z0-9]+", "-", path.stem)
    safe_stem = re.sub(r"-+", "-", safe_stem).strip("-")
    if not safe_stem:
        safe_stem = "document"
    return safe_stem + suffix


def prepare_download_artifact(source_path: str | Path, download_dir: str | Path = "/mnt/data") -> Path:
    source = Path(source_path)
    if not source.exists():
        raise FileNotFoundError(f"Rendered DOCX not found: {source}")
    if source.suffix.lower() != ".docx":
        raise ValueError(f"Expected .docx artifact, got: {source}")

    target_dir = Path(download_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / slugify_artifact_filename(source.name)

    if source.resolve() != target.resolve():
        shutil.copy2(source, target)

    validate_download_artifact(target)
    return target


def prepare_download_bundle(
    source_paths: list[str | Path],
    download_dir: str | Path = "/mnt/data",
    output_name: str | None = None,
) -> Path:
    if len(source_paths) < 2:
        raise ValueError("ZIP packaging requires at least two source documents.")

    sources = [Path(path) for path in source_paths]
    for source in sources:
        if not source.exists():
            raise FileNotFoundError(f"Rendered DOCX not found: {source}")
        if source.suffix.lower() != ".docx":
            raise ValueError(f"Expected .docx artifact, got: {source}")
        validate_download_artifact(source)

    target_dir = Path(download_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    bundle_name = slugify_zip_artifact_filename(output_name or infer_bundle_name(sources))
    target = target_dir / bundle_name

    used_names: set[str] = set()
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as package:
        for source in sources:
            archive_name = uniquify_archive_name(slugify_artifact_filename(source.name), used_names)
            package.write(source, arcname=archive_name)

    validate_zip_download_artifact(target, expected_count=len(sources))
    return target


def maybe_prepare_download_artifact(source_path: Path, download_dir: str | None) -> Path | None:
    if download_dir:
        return prepare_download_artifact(source_path, download_dir)
    default_dir = Path("/mnt/data")
    if default_dir.exists():
        return prepare_download_artifact(source_path, default_dir)
    return None


def validate_download_artifact(target: str | Path) -> None:
    artifact = Path(target)
    if not artifact.exists():
        raise FileNotFoundError(f"Download artifact was not created: {artifact}")
    if artifact.stat().st_size == 0:
        raise ValueError(f"Download artifact is empty: {artifact}")
    if artifact.suffix.lower() != ".docx":
        raise ValueError(f"Expected .docx download artifact, got: {artifact}")

    with zipfile.ZipFile(artifact, "r") as package:
        required = {"[Content_Types].xml", "word/document.xml"}
        missing = required - set(package.namelist())
        if missing:
            raise ValueError(f"Invalid DOCX package; missing: {sorted(missing)}")


def validate_zip_download_artifact(target: str | Path, expected_count: int | None = None) -> None:
    artifact = Path(target)
    if not artifact.exists():
        raise FileNotFoundError(f"Download artifact was not created: {artifact}")
    if artifact.stat().st_size == 0:
        raise ValueError(f"Download artifact is empty: {artifact}")
    if artifact.suffix.lower() != ".zip":
        raise ValueError(f"Expected .zip download artifact, got: {artifact}")

    with zipfile.ZipFile(artifact, "r") as package:
        names = [name for name in package.namelist() if not name.endswith("/")]
        if expected_count is not None and len(names) != expected_count:
            raise ValueError(f"Expected {expected_count} files in ZIP artifact, found {len(names)}")
        if not names:
            raise ValueError("ZIP artifact is empty.")
        for name in names:
            if Path(name).suffix.lower() != ".docx":
                raise ValueError(f"ZIP artifact contains a non-DOCX file: {name}")
            with package.open(name) as handle:
                data = handle.read()
            required = [b"[Content_Types].xml", b"word/document.xml"]
            if not all(token in data for token in required):
                raise ValueError(f"ZIP artifact member is not a valid DOCX package: {name}")


def infer_bundle_name(sources: list[Path]) -> str:
    stems = [source.stem for source in sources]
    if all("resume" in stem.lower() for stem in stems):
        return "resume-bundle.zip"
    return "document-bundle.zip"


def slugify_zip_artifact_filename(filename: str) -> str:
    path = Path(filename)
    stem = path.stem if path.suffix else filename
    safe_stem = re.sub(r"[^A-Za-z0-9]+", "-", stem)
    safe_stem = re.sub(r"-+", "-", safe_stem).strip("-")
    if not safe_stem:
        safe_stem = "documents"
    return safe_stem + ".zip"


def uniquify_archive_name(name: str, used_names: set[str]) -> str:
    candidate = name
    stem = Path(name).stem
    suffix = Path(name).suffix
    counter = 2
    while candidate in used_names:
        candidate = f"{stem}-{counter}{suffix}"
        counter += 1
    used_names.add(candidate)
    return candidate


def density_retry_order(start: str) -> list[str]:
    order = list(STYLE_PROFILES)
    index = order.index(start)
    return order[index:]


def density_attempts(density: str, doc_type: str, visual_qa_requested: bool) -> list[str]:
    if density != "auto":
        return density_retry_order(density) if visual_qa_requested else [density]
    if visual_qa_requested:
        return list(STYLE_PROFILES)
    return [estimated_fast_density(doc_type)]


def estimated_fast_density(doc_type: str) -> str:
    # Fast mode cannot know true Word pagination without a renderer, so it picks
    # a deterministic starting density and leaves exact page count to QA mode.
    return "normal" if doc_type == "cover_letter" else "compact"


def resolve_render_dir(render_dir: str | None, output: Path) -> Path:
    return Path(render_dir) if render_dir else output.parent / f"rendered-{output.stem}"


def resolve_effective_mode(
    requested_mode: str,
    cache_enabled: bool,
    cache_file: Path,
    doc_type: str,
    fingerprint: dict[str, Any],
) -> str:
    if requested_mode != "fast" or not cache_enabled or not cache_file.exists():
        return requested_mode
    cache = read_json_file(cache_file)
    record = cache.get(doc_type)
    if record and record.get("fingerprint") != fingerprint:
        return "qa"
    return requested_mode


def build_fingerprint(template: str | None, config: str | None = None) -> dict[str, Any]:
    renderer_path = Path(__file__).resolve()
    template_path = Path(template).resolve() if template else None
    config_path = Path(config).resolve() if config else Path(__file__).resolve().with_name(DEFAULT_CONFIG_FILE)
    return {
        "renderer": str(renderer_path),
        "renderer_hash": sha256_file(renderer_path),
        "template": str(template_path) if template_path else "",
        "template_hash": sha256_file(template_path) if template_path and template_path.exists() else "",
        "config": str(config_path) if config_path.exists() else "",
        "config_hash": sha256_file(config_path) if config_path.exists() else "",
        "validation_rules_version": VALIDATION_RULES_VERSION,
    }


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_json_file(path: Path) -> dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def update_validation_cache(cache_file: Path, doc_type: str, fingerprint: dict[str, Any], density: str) -> None:
    cache = read_json_file(cache_file)
    cache[doc_type] = {
        "fingerprint": fingerprint,
        "last_known_good_density": density,
    }
    cache_file.parent.mkdir(parents=True, exist_ok=True)
    cache_file.write_text(json.dumps(cache, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def build_qa_report(
    *,
    args: argparse.Namespace,
    model: dict[str, Any],
    output: Path,
    result,
    fingerprint: dict[str, Any],
    mode: str,
    density: str,
    page_count: int | None,
    visual_qa_performed: bool,
    visual_qa_reason: str,
    renderer_error: str = "",
    trim_report_path: Path | None = None,
    download_artifact: Path | None = None,
) -> dict[str, Any]:
    inspection = inspect_docx(output)
    status = result_status(result, renderer_error)
    if renderer_error and result.ok:
        status = RESULT_PASS_WITH_VISUAL_WARNING
    warnings = [issue.message for issue in result.issues if issue.severity == "warning"]
    errors = [issue.message for issue in result.issues if issue.severity == "error"]
    issue_codes = {issue.code for issue in result.issues}
    return {
        "source_file": str(Path(args.input)),
        "output_file": str(output),
        "download_artifact": str(download_artifact) if download_artifact else "",
        "document_type": model["type"],
        "template_used": args.template or "",
        "renderer": fingerprint["renderer"],
        "renderer_hash": fingerprint["renderer_hash"],
        "template_hash": fingerprint["template_hash"],
        "config_file": fingerprint.get("config", ""),
        "config_hash": fingerprint.get("config_hash", ""),
        "validation_rules_version": VALIDATION_RULES_VERSION,
        "mode": mode,
        "density_used": density,
        "max_pages": args.max_pages,
        "detected_pages": page_count,
        "visual_qa_performed": visual_qa_performed,
        "visual_qa_reason": visual_qa_reason,
        "hyperlinks_count": inspection.hyperlink_count,
        "raw_urls_visible": "RAW_URL" in issue_codes,
        "literal_markdown_bullets_visible": "LITERAL_BULLET" in issue_codes,
        "tables_present": inspection.table_count > 0,
        "columns_present": any(count > 1 for count in inspection.column_counts),
        "text_boxes_present": inspection.text_box_count > 0,
        "images_or_drawings_present": inspection.drawing_count > 0,
        "result": status,
        "warnings": warnings,
        "errors": errors,
        "renderer_error": renderer_error,
        "trim_report": str(trim_report_path) if trim_report_path else "",
        "word_count": estimate_word_count(model),
    }


def result_status(result, renderer_error: str = "") -> str:
    if renderer_error and not result.ok:
        return RESULT_FAIL_RENDERER
    if result.ok:
        return RESULT_PASS
    codes = {issue.code for issue in result.issues if issue.severity == "error"}
    if "PAGE_LIMIT" in codes:
        return RESULT_FAIL_PAGE_COUNT
    if any(code.startswith("HYPERLINK") for code in codes):
        return RESULT_FAIL_HYPERLINKS
    return RESULT_FAIL_STRUCTURE


def has_issue(result, code: str) -> bool:
    return any(issue.code == code for issue in result.issues)


def write_qa_report(report_path: Path | None, output: Path, report: dict[str, Any]) -> Path:
    path = report_path or output.with_suffix(".qa.json")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    return path


def write_trim_required_report(
    output: Path,
    model: dict[str, Any],
    result,
    density: str,
    max_pages: int | None,
) -> Path:
    path = output.with_suffix(".trim-required.txt")
    candidates = collect_trim_candidates(model)
    lines = [
        "Trim required: rendered output exceeded the page target.",
        f"Output: {output}",
        f"Density: {density}",
        f"Max pages: {max_pages}",
        "",
        "No content was removed automatically.",
        "",
        "Validation errors:",
    ]
    lines.extend(f"- {issue.code}: {issue.message}" for issue in result.issues if issue.severity == "error")
    lines.extend(["", "Candidate cuts for manual review:"])
    if candidates:
        for index, candidate in enumerate(candidates[:12], start=1):
            lines.append(f"{index}. {candidate}")
    else:
        lines.append("- No obvious bullet-level trim candidates were detected.")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def collect_trim_candidates(model: dict[str, Any]) -> list[str]:
    candidates: list[str] = []
    for section in reversed(model.get("sections", [])):
        section_heading = section.get("heading", "Section")
        for item in reversed(section.get("items", [])):
            item_heading = item.get("heading") or item.get("components", {}).get("role", "")
            for block in reversed(item.get("blocks", [])):
                candidates.extend(format_trim_block(section_heading, item_heading, block))
        for block in reversed(section.get("blocks", [])):
            candidates.extend(format_trim_block(section_heading, "", block))
    return candidates


def format_trim_block(section_heading: str, item_heading: str, block: dict[str, Any]) -> list[str]:
    prefix = f"{section_heading}"
    if item_heading:
        prefix += f" / {item_heading}"
    if block.get("kind") == "bullet_list":
        return [f"{prefix}: {plain_run_text(item)[:160]}" for item in block.get("items", [])]
    if block.get("kind") == "paragraph":
        text = plain_run_text(block.get("runs", []))
        return [f"{prefix}: {text[:160]}"] if text else []
    return []


def plain_run_text(runs: list[dict[str, Any]]) -> str:
    return "".join(run.get("text", "") for run in runs).strip()


if __name__ == "__main__":
    raise SystemExit(main())

from pathlib import Path
import argparse, json, shutil

SAMPLE_MD = """# Business Case: {initiative_name}

**Prepared for:** {prepared_for}  
**Prepared by:** Business Case System  
**Version:** Generated sample

## Executive Summary

{problem}

Recommended decision: {decision}

## Problem Statement and Business Need

Root-cause hypothesis: {root_cause}

## Financial Analysis

| Metric | Value |
|---|---:|
| One-time cost | ${one_time_cost:,.0f} |
| Annual recurring cost | ${annual_recurring_cost:,.0f} |
| Annual hard benefit | ${annual_hard_benefit:,.0f} |
| Approximate payback | {payback_months:.1f} months |

## Recommendation and Call to Action

Approve the next governed phase only after validating the operating baseline, financial assumptions, and implementation dependencies.
"""

def html_wrap(md):
    import html, re
    out=[]
    for line in md.splitlines():
        if line.startswith('# '): out.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith('## '): out.append(f"<h2>{html.escape(line[3:])}</h2>")
        elif line.startswith('|'): out.append(f"<pre>{html.escape(line)}</pre>")
        elif line.strip(): out.append(f"<p>{html.escape(line)}</p>")
    return '<!doctype html><html><head><meta charset="utf-8"><title>Business Case</title></head><body>' + '\n'.join(out) + '</body></html>'

def make_docx(md, path):
    from docx import Document
    doc=Document()
    for line in md.splitlines():
        if line.startswith('# '): doc.add_heading(line[2:], 0)
        elif line.startswith('## '): doc.add_heading(line[3:], 1)
        elif line.strip() and not line.startswith('|'): doc.add_paragraph(line)
    doc.save(path)

def build(args):
    root = Path(__file__).resolve().parents[3]
    sample_path = root / 'examples' / 'sample-data' / 'sample_business_case_input.json'
    data = json.loads(sample_path.read_text(encoding='utf-8')) if args.sample else json.loads(Path(args.input).read_text(encoding='utf-8'))
    fin = data.get('financials', {})
    md = SAMPLE_MD.format(
        initiative_name=data.get('initiative_name','Untitled Business Case'),
        prepared_for=data.get('prepared_for','Decision-maker'),
        problem=data.get('problem','Problem not supplied.'),
        root_cause=data.get('root_cause','Root-cause hypothesis not supplied.'),
        decision=data.get('decision','Decision not supplied.'),
        one_time_cost=fin.get('one_time_cost',0),
        annual_recurring_cost=fin.get('annual_recurring_cost',0),
        annual_hard_benefit=fin.get('annual_hard_benefit',0),
        payback_months=fin.get('payback_months',0),
    )
    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    formats = args.formats.split(',')
    if 'md' in formats: (out/'business_case.md').write_text(md, encoding='utf-8')
    if 'html' in formats: (out/'business_case.html').write_text(html_wrap(md), encoding='utf-8')
    if 'docx' in formats: make_docx(md, out/'business_case.docx')
    print(f"Generated outputs in {out}")

def main():
    parser = argparse.ArgumentParser(prog='business-case-system')
    sub = parser.add_subparsers(dest='command', required=True)
    b = sub.add_parser('build')
    b.add_argument('--sample', action='store_true')
    b.add_argument('--input')
    b.add_argument('--out', default='out')
    b.add_argument('--formats', default='md,html,docx')
    b.set_defaults(func=build)
    args = parser.parse_args()
    args.func(args)

if __name__ == '__main__':
    main()
